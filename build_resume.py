# -*- coding: utf-8 -*-
"""
Renders Kshitiz Yadav's one-page ATS resume as a .docx.

DATA holds all resume content as a plain dict so a tailored variant (built by hand when
processing a "Tailor My Resume" request -- see CLAUDE.md) can pass a modified copy
(reworded/reordered bullets per a JD) through the same render() without duplicating the
docx layout code. keyword_coverage_score() below scores a tailored variant against a JD.

Structure (2026-08-05 revamp): the resume now shows the REAL 3-role progression at PlanetSpark
(Senior Product Analyst -> Assistant Product Manager -> Product Manager) as separate dated
sub-entries under one company, a compact honest metrics strip under the header, and Core
Competencies placed directly BELOW the summary (so recruiters see transferable positioning
before the EdTech employer context). Email + LinkedIn render as real clickable hyperlinks.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x5A, 0xAC)
GREY = RGBColor(0x44, 0x44, 0x44)
DGREY = RGBColor(0x33, 0x33, 0x33)
FONT = "Calibri"

# ---- style presets ----
# render() takes a `style` dict controlling the knobs that drive vertical page-fill.
# fit_to_page() (below) picks the loosest preset that keeps the resume on exactly ONE page
# while filling it cleanly (~93-98%), so there's never a big empty gap at the bottom and it
# never spills onto a second page -- for any content volume (base or tailored variant).
# body = base body-font pt; head/name/company scale proportionally off it. line = line spacing.
# gap = space_after (pt) on body paragraphs & bullets; sec_before/sec_after = header spacing;
# sub_before = space before a bold sub-heading.
STYLE_DEFAULT = {
    "body": 9.5, "line": 1.0, "gap": 1.0,
    "sec_before": 5.0, "sec_after": 2.0, "sub_before": 3.0,
}

# Ordered tight -> loose. fit_to_page() walks these and keeps the fullest one that still fits
# on a single page. Fonts stay in the ATS-safe / readable 9.5-11pt band throughout. Steps are
# intentionally fine-grained (font AND line/gap increments interleaved) so a resume whose content
# sits between two font sizes -- where the bigger font overflows to 2 pages but the smaller one
# under-fills -- can still be nudged into the ~93-98% band via line spacing rather than falling
# back to a visibly short page.
STYLE_PRESETS = [
    {"body": 9.5,  "line": 1.00, "gap": 1.0, "sec_before": 5.0, "sec_after": 2.0, "sub_before": 3.0},
    {"body": 9.5,  "line": 1.06, "gap": 2.0, "sec_before": 6.0, "sec_after": 2.5, "sub_before": 3.5},
    {"body": 9.5,  "line": 1.11, "gap": 2.5, "sec_before": 6.5, "sec_after": 2.5, "sub_before": 3.5},
    {"body": 9.5,  "line": 1.16, "gap": 3.0, "sec_before": 7.0, "sec_after": 3.0, "sub_before": 4.0},
    {"body": 9.75, "line": 1.14, "gap": 3.0, "sec_before": 7.0, "sec_after": 3.0, "sub_before": 4.0},
    {"body": 9.5,  "line": 1.22, "gap": 3.5, "sec_before": 7.5, "sec_after": 3.0, "sub_before": 4.0},
    {"body": 10.0, "line": 1.06, "gap": 2.0, "sec_before": 6.0, "sec_after": 2.5, "sub_before": 3.5},
    {"body": 10.0, "line": 1.12, "gap": 3.0, "sec_before": 7.0, "sec_after": 3.0, "sub_before": 4.0},
    {"body": 10.0, "line": 1.18, "gap": 3.5, "sec_before": 7.5, "sec_after": 3.5, "sub_before": 4.5},
    {"body": 10.5, "line": 1.10, "gap": 3.0, "sec_before": 7.0, "sec_after": 3.0, "sub_before": 4.0},
    {"body": 10.5, "line": 1.16, "gap": 4.0, "sec_before": 8.0, "sec_after": 3.5, "sub_before": 4.5},
    {"body": 11.0, "line": 1.15, "gap": 4.0, "sec_before": 8.0, "sec_after": 4.0, "sub_before": 5.0},
    {"body": 11.0, "line": 1.22, "gap": 5.0, "sec_before": 9.0, "sec_after": 4.0, "sub_before": 5.5},
]

DATA = {
    "name": "KSHITIZ YADAV",
    "tagline": "Product Manager  |  Growth, Monetization & Payments  |  Consumer, Marketplace & Platform Products  |  SQL, Analytics & APIs",
    "contact": "Gurugram, India  •  +91 8756972501  •  kshitizyadav788@gmail.com  •  linkedin.com/in/kshitizyadav",
    # links rendered as real clickable hyperlinks (see render()); keep these in sync with `contact`.
    "email": "kshitizyadav788@gmail.com",
    "linkedin": "linkedin.com/in/kshitizyadav",
    # Compact, honest headline metrics strip shown under the header. Real numbers only.
    "metrics": "20% Revenue Growth   •   25% ARPU Expansion (₹35K → ₹45K)   •   ₹70L Renewals in a Single Month   •   80% Lower Ops Workload",
    # Keep this SHORT (2-3 lines): positioning only -- who he is, what he owns, how he works, one
    # or two headline numbers. Do NOT restate the metrics that already appear in the bullets below;
    # that duplication was the reason this was rewritten (see CLAUDE.md "Master-resume conventions").
    "summary": (
        "Product Manager with 5+ years of product experience building growth, monetization, payments, and "
        "workflow-automation products across consumer apps, a two-sided marketplace, CRM, and internal platforms "
        "serving 150,000+ enrolled users. Progressed from Senior Product Analyst to Assistant Product Manager to "
        "Product Manager — owning discovery, PRDs, prioritization, launch, and iteration with engineering, design, "
        "and business teams. Delivered 20% revenue growth, 25% ARPU expansion, ₹70L in product-led renewals in a "
        "single month, and 80% lower operational workload, using SQL, Power BI, and AI as force-multipliers."
    ),
    # Placed directly below the summary in render(). Transferable-first ordering for an industry switch.
    "competencies": "Product Strategy & Roadmapping · Product Discovery & User Research · PRDs & Prioritization · Growth & Monetization · Pricing & Packaging · Retention & Lifecycle · Funnel & Conversion Optimization · Payments & API Integrations · Workflow Automation · Product Analytics & Insights · GTM & Product Launch · Marketplace & Platform Products · Unit Economics · Cross-Functional Leadership",
    "company": "PlanetSpark — Gurugram",
    "company_intro": "A B2C consumer marketplace (EdTech) connecting 150,000+ enrolled users with teachers via a shared LMS — spanning consumer app, two-sided platform, payments, CRM, renewals, and lifecycle products, with 30,000+ daily active users. Progressed across three roles while owning an expanding product surface.",
    "roles": [
        {"title": "Product Manager", "dates": "Sep 2023 – Present", "bullets": [
            "Own the full product suite — consumer app, student & teacher LMS, advisor CRM, renewal engine, and in-house AI layer — for a platform with 150,000+ enrolled users and 30,000+ daily active users, leading a cross-functional squad of 8 engineers and 2 QA from discovery to launch.",
            "Designed an in-product renewal engine (nudges, early-bird access, LMS free-trial classes, teacher-initiated renewals) that reduced assisted-renewal cost and generated ₹70L in product-led renewals in a single month (Jan 2025).",
            "Increased the share of enrolled users rating the product 5 out of 5 from 10% to 25% via onboarding and retention work: seamless onboarding, teacher PTMs and feedback, a one-page student-detail UI, a loyalty program, and the PSpark in-app wallet.",
            "Re-architected the sign-up → demo → enrolment funnel — single-step OTP, a preference-capture ranking model, and 6 API-integrated payment gateways with pre-filled links and auto-cleared approvals — enabling self-serve, plan-based purchase with no manual intervention.",
            "Built an advisor-wise P&L tracker (weekly Net Revenue − Fixed, Refund, Marketing & Sales cost) whose visibility drove profitability interventions — profitable advisors rose from 50 to 100 in one month and to 125 over the next two.",
            "Shipped a production GenAI learning layer (age-appropriate speech/text responses) and a RAG-based support chatbot — cutting first-response time from 4 minutes to under 10 seconds and monthly escalations from ~300 to under 50 within three months.",
            "Ship minor and mid-level product features using AI-assisted development (Claude Code) — scoping, implementing, and testing changes, and raising merge requests for engineering review.",
        ]},
        {"title": "Assistant Product Manager", "dates": "Sep 2022 – Aug 2023", "bullets": [
            "Built the LPP (Learn, Practice, Perform) module — customized 1:1 sessions, group activities, and performance tasks — driving 20% revenue growth, 25% ARPU expansion (₹35K → ₹45K), and 50% higher revenue per class (₹600 → ₹900).",
            "Rebuilt the teacher LMS with real-time payout automation, accurate payout summaries, and consolidated CRM and class details — lifting teacher NPS from −20 to +40.",
            "Automated enrolment verification and incentive flows to 100% coverage and revamped the advisor CRM (lead navigation, calling, revenue & target tracking) — cutting operations workload 80% and saving ₹1.6L/month.",
            "Launched organic acquisition loops surfacing student progress (Sparkline, Word Wisdom, practice classes, workshops) across social channels, bringing 2,000+ organic leads/month at a 7% conversion rate.",
        ]},
        {"title": "Senior Product Analyst", "dates": "Mar 2021 – Aug 2022", "bullets": [
            "Developed SQL and Power BI reporting across growth, revenue, and operations; improved weekly revenue accuracy and created a data foundation for product and commercial decisions.",
            "Implemented customer-verification controls that reduced invalid or fraudulent transactions from 10% to 1% while minimizing manual intervention.",
        ]},
    ],
    "skills": [
        ("Data & Analytics", "SQL, GA4, Power BI, MS Excel (Advanced), Python (basic)"),
        ("Product & Delivery", "JIRA, Confluence, Agile/Scrum, RICE prioritization, Figma"),
        ("AI & Integration", "OpenAI / Gemini / Claude APIs, Claude Code, Prompt Engineering, RAG, REST APIs, Payment Integrations"),
    ],
    "education": [
        ("BBA — Chandigarh University, Chandigarh", "2018 – 2021"),
    ],
    "certifications": "Product Management Certification — Udemy (2024)    •    JIRA Certification — Udemy (2024)",
}


def render(data, out_path, style=None):
    s = dict(STYLE_DEFAULT)
    if style:
        s.update(style)
    body = s["body"]; line = s["line"]; gap = s["gap"]
    sec_before = s["sec_before"]; sec_after = s["sec_after"]; sub_before = s["sub_before"]
    # secondary fonts scale off body so a bump reads proportionally, staying ATS-readable
    head = body + 1.5; name_sz = body + 7.5; company_sz = body + 1.0
    tagline_sz = body; contact_sz = body - 0.5

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(body)
    normal.paragraph_format.space_after = Pt(gap)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = line

    try:
        lb = doc.styles["List Bullet"]
        lb.font.name = FONT; lb.font.size = Pt(body)
        lb.paragraph_format.space_after = Pt(gap)
        lb.paragraph_format.space_before = Pt(0)
        lb.paragraph_format.line_spacing = line
    except KeyError:
        pass

    sec = doc.sections[0]
    sec.page_width = Mm(210); sec.page_height = Mm(297)        # A4
    sec.top_margin = Mm(9); sec.bottom_margin = Mm(9)
    sec.left_margin = Mm(12); sec.right_margin = Mm(12)

    def set_bottom_border(p):
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '2E5AAC')
        pbdr.append(bottom); pPr.append(pbdr)

    def header_line(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(sec_before); p.paragraph_format.space_after = Pt(sec_after)
        r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(head)
        r.font.color.rgb = NAVY; r.font.name = FONT
        set_bottom_border(p)

    def add_run(p, text, bold=False, italic=False, size=None, color=None):
        r = p.add_run(text); r.bold = bold; r.italic = italic
        r.font.size = Pt(size if size is not None else body); r.font.name = FONT
        if color is not None: r.font.color.rgb = color
        return r

    def add_hyperlink(p, url, text, size, color):
        r_id = p.part.relate_to(
            url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True)
        link = OxmlElement('w:hyperlink'); link.set(qn('r:id'), r_id)
        run = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
        rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'), FONT); rf.set(qn('w:hAnsi'), FONT); rPr.append(rf)
        c = OxmlElement('w:color'); c.set(qn('w:val'), '%02X%02X%02X' % (color[0], color[1], color[2])); rPr.append(c)
        szv = OxmlElement('w:sz'); szv.set(qn('w:val'), str(int(size * 2))); rPr.append(szv)
        run.append(rPr)
        t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text; run.append(t)
        link.append(run); p._p.append(link)

    def bullet(parts):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(gap); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = line
        if isinstance(parts, str): add_run(p, parts)
        else:
            for txt, kw in parts: add_run(p, txt, **kw)

    def right_tab(p):
        cw = sec.page_width - sec.left_margin - sec.right_margin
        p.paragraph_format.tab_stops.add_tab_stop(cw, WD_TAB_ALIGNMENT.RIGHT)

    # -------- HEADER --------
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0)
    add_run(p, data["name"], bold=True, size=name_sz, color=NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0)
    add_run(p, data["tagline"], size=tagline_sz, color=GREY)
    # contact line with clickable email + LinkedIn
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(1)
    email = data.get("email"); linkedin = data.get("linkedin")
    if email and linkedin:
        add_run(p, "Gurugram, India  •  +91 8756972501  •  ", size=contact_sz, color=GREY)
        add_hyperlink(p, "mailto:" + email, email, contact_sz, GREY)
        add_run(p, "  •  ", size=contact_sz, color=GREY)
        add_hyperlink(p, "https://" + linkedin, linkedin, contact_sz, GREY)
    else:
        add_run(p, data["contact"], size=contact_sz, color=GREY)
    # metrics strip
    if data.get("metrics"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(gap)
        add_run(p, data["metrics"], bold=True, size=contact_sz, color=BLUE)

    # -------- SUMMARY --------
    header_line("Professional Summary")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(gap)
    add_run(p, data["summary"])

    # -------- CORE COMPETENCIES (directly below summary) --------
    header_line("Core Competencies")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(gap)
    add_run(p, data["competencies"])

    # -------- EXPERIENCE (one company, 3 dated role sub-entries) --------
    header_line("Professional Experience")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    add_run(p, data["company"], bold=True, size=company_sz)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(gap)
    add_run(p, data["company_intro"], italic=True, color=DGREY)
    for role in data["roles"]:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(sub_before)
        p.paragraph_format.space_after = Pt(0); right_tab(p)
        add_run(p, role["title"], bold=True, italic=True, color=BLUE)
        add_run(p, "\t" + role["dates"], italic=True, color=GREY)
        for b in role["bullets"]:
            bullet(b)

    # -------- TECHNICAL SKILLS --------
    header_line("Technical Skills & Tools")
    for label, text in data["skills"]:
        bullet([(label + ": ", {"bold": True}), (text, {})])

    # -------- EDUCATION & CERTIFICATIONS --------
    header_line("Education & Certifications")
    for school, dates in data["education"]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(gap); right_tab(p)
        add_run(p, school, bold=True); add_run(p, "\t" + dates, color=GREY)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(gap)
    add_run(p, data["certifications"], color=GREY)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path


def measure_pdf(pdf_path):
    """Return (page_count, fill_pct) for a rendered resume PDF. fill_pct is how far the lowest
    text baseline reaches down the first page as a % of page height -- our proxy for 'how full'."""
    from pypdf import PdfReader
    r = PdfReader(pdf_path)
    pages = len(r.pages)
    pg = r.pages[0]
    h = float(pg.mediabox.height)
    lowest = [h]
    def visitor(text, cm, tm, font, size):
        if text.strip():
            lowest[0] = min(lowest[0], tm[5])
    pg.extract_text(visitor_text=visitor)
    fill_pct = round(100 * (h - lowest[0]) / h, 1)
    return pages, fill_pct


def _render_to_pdf(data, docx_path, style):
    """Render docx with a style, convert to PDF via LibreOffice, return (pages, fill_pct, pdf_path)."""
    import subprocess, shutil
    render(data, docx_path, style=style)
    outdir = os.path.dirname(os.path.abspath(docx_path))
    soffice = shutil.which("soffice") or "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
                   check=True, capture_output=True)
    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    pages, fill = measure_pdf(pdf_path)
    return pages, fill, pdf_path


def fit_to_page(data, out_path, target_lo=93.0, target_hi=98.0, verbose=True):
    """Render `data` at the loosest STYLE preset that still fits on exactly ONE page and fills
    it cleanly (fill% ideally in [target_lo, target_hi]). Requires LibreOffice (soffice) for the
    real page measurement -- there is no pure-Python way to know rendered page count/fill.

    Strategy: walk presets tight -> loose. Track the fullest single-page result seen. Stop early
    once a single-page result lands in the target band. Never returns a 2-page layout if any
    single-page option exists. Returns (chosen_style, pages, fill_pct, pdf_path).

    If even the tightest preset spills to 2 pages, the CONTENT is too long -- trim bullets/summary
    (see CLAUDE.md), don't just shrink further. If the loosest preset still under-fills, the
    content is too thin -- add a real bullet from experience_bank.md rather than bloating spacing."""
    best = None  # (fill, style, pages, pdf) among single-page renders
    for st in STYLE_PRESETS:
        pages, fill, pdf = _render_to_pdf(data, out_path, st)
        if verbose:
            print(f"  body={st['body']:>4} line={st['line']:.2f} gap={st['gap']:.0f} -> pages={pages} fill={fill}%")
        if pages == 1:
            if best is None or fill > best[0]:
                best = (fill, st, pages, pdf)
            if target_lo <= fill <= target_hi:
                break
    if best is None:
        # every preset overflowed: re-render tightest so output is at least the least-bad, and warn
        st = STYLE_PRESETS[0]
        pages, fill, pdf = _render_to_pdf(data, out_path, st)
        print(f"  WARNING: content too long for one page even at tightest preset (pages={pages}). "
              f"Trim content -- see CLAUDE.md.")
        return st, pages, fill, pdf
    fill, st, pages, pdf = best
    if fill < target_lo:
        print(f"  NOTE: best fill only {fill}% (< {target_lo}%). Content is thin -- consider adding "
              f"a real bullet from experience_bank.md rather than relying on spacing.")
    # ensure the on-disk docx matches the chosen style (last render may have been a different preset)
    render(data, out_path, style=st)
    _render_to_pdf(data, out_path, st)
    return st, pages, fill, pdf


def keyword_coverage_score(data, keywords):
    """Deterministic, honest match score: % of an EXPLICIT curated keyword/must-have list
    (extracted by hand from the JD while tailoring -- see CLAUDE.md step (b)) that actually
    appears in the tailored resume text. This is deliberately NOT raw full-JD-text word
    overlap -- a naive version of that scored a real tailored resume at 33% because JDs are
    full of generic connector words ("translate", "ensuring", "deliver") no resume would ever
    contain even when the underlying skill is genuinely covered. Real ATS systems and
    recruiters key off specific skill/tool/domain terms, not prose glue -- so should this.

    A multi-word keyword phrase matches if ALL of its significant words appear somewhere in the
    resume (not necessarily adjacent) -- exact contiguous substring matching was tried first and
    produced false negatives: "product development" didn't match "Product Ideation & Development"
    (words not adjacent), "market analysis" didn't match "market research and competitor analysis"
    (same concept, different word order/filler). Single-word keywords still match as substrings.

    `keywords` must be the actual must-have terms judged from reading the JD, not padded to
    inflate the score; the resulting number is only honest if that list is honest. Returns
    (score, matched, missing) so the gaps are visible, not just a bare percentage."""
    _filler = {"and", "the", "of", "a", "an", "for", "to", "in", "&"}
    bullets = " ".join(b for r in data["roles"] for b in r["bullets"])
    resume_text = " ".join([
        data.get("summary", ""), data.get("competencies", ""), data.get("tagline", ""),
        data.get("company_intro", ""), data.get("metrics", ""),
        " ".join(t for _, t in data.get("skills", [])),
        bullets,
    ]).lower()

    def is_match(keyword):
        words = [w for w in keyword.lower().split() if w not in _filler]
        return all(w in resume_text for w in words) if words else False

    matched = [k for k in keywords if is_match(k)]
    missing = [k for k in keywords if not is_match(k)]
    score = round(100 * len(matched) / len(keywords)) if keywords else None
    return score, matched, missing


if __name__ == "__main__":
    _here = os.path.dirname(os.path.abspath(__file__))
    out = os.path.join(_here, "resumes", "Kshitiz Yadav - Product Manager Resume.docx")
    print("Auto-fitting base resume to one full page...")
    style, pages, fill, pdf = fit_to_page(DATA, out)
    print(f"saved: {out}")
    print(f"       {pdf}")
    print(f"final: pages={pages}, fill={fill}%, style body={style['body']}pt line={style['line']}")
