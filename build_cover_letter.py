# -*- coding: utf-8 -*-
"""
Renders a one-page cover letter as a .docx, matching the resume's visual style
(same fonts/colors from build_resume.py) so the two documents look like a set.

Usage: write the letter body as a list of paragraph strings, then call
render(paragraphs, out_path, company, role).
"""
import os
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import build_resume as br  # reuse NAVY/GREY/FONT and contact details

BODY = 11.0
NAME = 15.0


def render(paragraphs, out_path, company, role):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = br.FONT
    normal.font.size = Pt(BODY)
    normal.paragraph_format.space_after = Pt(10)
    normal.paragraph_format.line_spacing = 1.15

    sec = doc.sections[0]
    sec.page_width = Mm(210); sec.page_height = Mm(297)
    sec.top_margin = Mm(20); sec.bottom_margin = Mm(20)
    sec.left_margin = Mm(22); sec.right_margin = Mm(22)

    def run(p, text, bold=False, size=BODY, color=None):
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = br.FONT
        if color is not None: r.font.color.rgb = color
        return r

    p = doc.add_paragraph(); run(p, "KSHITIZ YADAV", bold=True, size=NAME, color=br.NAVY)
    p = doc.add_paragraph()
    run(p, "Gurugram, India  •  +91 8756972501  •  kshitizyadav788@gmail.com  •  linkedin.com/in/kshitizyadav", size=9.5, color=br.GREY)

    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14)
    import datetime
    run(p, datetime.date.today().strftime("%B %d, %Y"), color=br.GREY)

    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10)
    run(p, f"Re: {role}, {company}", bold=True)

    for para in paragraphs:
        p = doc.add_paragraph()
        run(p, para)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path
